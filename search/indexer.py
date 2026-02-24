"""
BFS streaming indexer with checkpointing and robust PDF pipeline.
"""

import os
import json
import time
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Set, Tuple
import hashlib

from .config import get_config
from .storage import create_storage
from .ocr import get_ocr_extractor
from .types import Chunk, FrontierState, IndexStats
from .ids import file_id, chunk_id, generate_file_sha256, get_file_stats
from .model_loader import get_embedding_model

logger = logging.getLogger(__name__)

class BFSIndexer:
    """BFS streaming indexer with checkpointing."""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or get_config()
        self.qdrant, self.catalog = create_storage(self.config)
        self.frontier_path = Path(self.config["paths"]["frontier"])
        self.max_items = self.config["index"].get("max_items", 1000)
        self.exclude_patterns = self.config["index"]["exclude_patterns"]
        self.allow_exts = set(self.config["index"]["allow_exts"])
        self.ocr_enabled = self.config["index"].get("ocr_enabled", False)
        self.ocr_backend = self.config["index"].get("ocr_backend", "tesseract")
        self.ocr_only_for_images = self.config["index"].get("ocr_only_for_images", True)
        self.ocr_paths = self.config["index"].get("ocr_paths") or []
        if self.ocr_enabled:
            self.allow_exts |= {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"}
        self._image_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"}
        self._text_exts = {".txt", ".md", ".markdown", ".pdf", ".docx", ".html", ".htm", ".rtf"}
        self.max_pdf_pages = self.config["index"]["max_pdf_pages"]
        self.extraction_timeout = self.config["index"]["extraction_timeout"]
        
        # Stats tracking
        self.stats = IndexStats()
        self._closed = False
        # Optimization: accumulate chunks before embedding (batch across files)
        self._embed_buffer: List[Chunk] = []
        self._embed_accumulate_batch = self.config["index"].get("embed_accumulate_batch", 2048)
    
    def close(self):
        """Close all resources (database connections, etc.)."""
        if self._closed:
            return
        
        try:
            if hasattr(self, 'catalog') and self.catalog:
                self.catalog.close()
            logger.debug("BFSIndexer resources closed")
        except Exception as e:
            logger.warning(f"Error closing BFSIndexer resources: {e}")
        finally:
            self._closed = True
    
    def __enter__(self):
        """Context manager entry - allows 'with BFSIndexer(...)' usage."""
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit - automatically closes resources."""
        self.close()
        return False  # Don't suppress exceptions
    
    def __del__(self):
        """Cleanup on deletion - ensures resources are closed."""
        if not self._closed:
            try:
                self.close()
            except Exception:
                pass  # Ignore errors during cleanup
    
    def run_bfs_slice(self, roots: List[str], max_items: int = None) -> IndexStats:
        """Run one BFS slice with checkpointing."""
        max_items = max_items or self.max_items
        
        # Load or create frontier
        frontier = self._load_frontier()
        
        # Add roots to frontier if empty
        if not frontier.queue:
            for root in roots:
                if Path(root).exists():
                    frontier.queue.append(root)
                    logger.info(f"Added root to frontier: {root}")
        
        # Process one level
        processed_count = 0
        current_level = []
        
        # Get current level items
        while frontier.queue and processed_count < max_items:
            current_level.append(frontier.queue.pop(0))
            processed_count += 1
        
        logger.info(f"Processing {len(current_level)} items from frontier")
        
        # Prioritize text files (fast) over images (slow OCR) for better perceived indexing speed
        current_level = self._sort_for_processing(current_level)
        
        # Process current level
        for item_path in current_level:
            try:
                self._process_item(item_path, frontier)
            except Exception as e:
                logger.error(f"Error processing {item_path}: {e}")
                frontier.errors.append(f"{item_path}: {str(e)}")
                self.stats.errors += 1
        
        # Flush any remaining chunks in embed buffer
        self._flush_embed_buffer()
        
        # Save frontier state
        self._save_frontier(frontier)
        
        return self.stats
    
    def _process_item(self, item_path: str, frontier: FrontierState):
        """Process a single file or directory."""
        path = Path(item_path)
        
        if not path.exists():
            logger.warning(f"Path does not exist: {item_path}")
            return
        
        # Check if already processed
        device_inode = (path.stat().st_dev, path.stat().st_ino)
        if item_path in frontier.seen and frontier.seen[item_path] == device_inode:
            logger.debug(f"Skipping already processed: {item_path}")
            return
        
        if path.is_file():
            self._process_file(item_path)
            frontier.processed_files += 1
            self.stats.files_processed += 1
        elif path.is_dir():
            self._process_directory(item_path, frontier)
            frontier.processed_dirs += 1
        
        # Mark as seen
        frontier.seen[item_path] = device_inode
    
    def _sort_for_processing(self, items: List[str]) -> List[str]:
        """Prioritize text files (fast) over images (slow OCR) for better indexing throughput."""
        def _priority(item: str) -> int:
            p = Path(item)
            if not p.exists() or not p.is_file():
                return 1  # dirs and missing: middle
            ext = p.suffix.lower()
            if ext in self._text_exts:
                return 0  # text first
            if ext in self._image_exts:
                return 2  # images last (OCR is slow)
            return 1
        return sorted(items, key=_priority)
    
    def _image_in_ocr_paths(self, file_path: str) -> bool:
        """Check if image path matches ocr_paths (when set). Empty ocr_paths = all match."""
        if not self.ocr_paths:
            return True
        path_norm = str(Path(file_path).resolve()).lower()
        return any(part.lower() in path_norm for part in self.ocr_paths)
    
    def _process_file(self, file_path: str):
        """Process a single file."""
        path = Path(file_path)
        
        logger.info(f"Processing file: {file_path}")
        
        # Check file extension
        if path.suffix.lower() not in self.allow_exts:
            logger.info(f"Skipping unsupported file (extension {path.suffix.lower()}): {file_path}")
            self.stats.files_skipped += 1
            return
        
        # OCR only for images in ocr_paths when configured (skip images elsewhere for faster indexing)
        if (path.suffix.lower() in self._image_exts and 
                self.ocr_enabled and 
                not self._image_in_ocr_paths(file_path)):
            logger.debug(f"Skipping image outside ocr_paths: {file_path}")
            self.stats.files_skipped += 1
            return
        
        # Check exclude patterns
        if self._should_exclude(file_path):
            logger.info(f"Excluded by pattern: {file_path}")
            self.stats.files_skipped += 1
            return
        
        # Get file stats
        stats = get_file_stats(file_path)
        if not stats:
            logger.warning(f"Could not get stats for: {file_path}")
            return
        
        # Generate file ID and check if unchanged
        fid = file_id(file_path, stats["mtime"], stats["size"])
        
        # Check if file already exists and is unchanged
        existing_sha256 = self._get_existing_sha256(fid)
        if existing_sha256:
            current_sha256 = generate_file_sha256(file_path)
            if current_sha256 == existing_sha256:
                logger.debug(f"File unchanged, skipping: {file_path}")
                self.stats.files_skipped += 1
                return
        
        # Extract text
        text = self._extract_text(file_path)
        if not text:
            logger.warning(f"No text extracted from: {file_path}")
            self.stats.files_skipped += 1
            return
        
        # Generate new SHA256
        new_sha256 = hashlib.sha256(text.encode()).hexdigest()
        
        # Chunk text (before database operations)
        chunks = self._chunk_text(text, file_path, fid)
        
        # All database operations in one transaction (all-or-nothing)
        try:
            with self.catalog.transaction():
                # Step 1: Update file metadata
                self.catalog.upsert_file(
                    str(file_path), stats["size"], stats["mtime"], new_sha256,
                    in_transaction=True
                )
                
                # Step 2: Insert chunks into catalog
                self.catalog.insert_chunks(fid, chunks, in_transaction=True)
                
                # Step 3: Insert into FTS (all chunks in same transaction)
                for chunk in chunks:
                    self.catalog.fts_insert(
                        chunk.chunk_id, chunk.text, chunk.path,
                        in_transaction=True
                    )
                # All database operations commit together here, or all rollback on error
        except Exception as e:
            logger.error(f"Database operations failed for {file_path}: {e}")
            raise
        
        # Step 4: Add chunks to embed buffer; flush when batch is full
        self._embed_buffer.extend(chunks)
        if len(self._embed_buffer) >= self._embed_accumulate_batch:
            self._flush_embed_buffer()
        
        self.stats.chunks_created += len(chunks)
        logger.info(f"Processed file: {file_path} ({len(chunks)} chunks)")
    
    def _process_directory(self, dir_path: str, frontier: FrontierState):
        """Process directory and add children to frontier."""
        try:
            path = Path(dir_path)
            
            # Check exclude patterns
            if self._should_exclude(dir_path):
                logger.debug(f"Excluded directory: {dir_path}")
                return
            
            # Add children to frontier
            try:
                children_added = 0
                for child in path.iterdir():
                    child_str = str(child)
                    
                    # Skip hidden files/directories
                    if child.name.startswith('.'):
                        continue
                    
                    # Skip if should be excluded
                    if self._should_exclude(child_str):
                        continue
                    
                    frontier.queue.append(child_str)
                    children_added += 1
                
                logger.info(f"Added {children_added} children from directory: {dir_path}")
                    
            except PermissionError:
                logger.warning(f"Permission denied accessing: {dir_path}")
            except OSError as e:
                logger.warning(f"Error reading directory {dir_path}: {e}")
                
        except Exception as e:
            logger.error(f"Error processing directory {dir_path}: {e}")
    
    def _extract_text(self, file_path: str) -> Optional[str]:
        """Extract text from file with robust pipeline."""
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        try:
            if suffix == '.txt':
                return self._extract_txt(file_path)
            elif suffix == '.md':
                return self._extract_md(file_path)
            elif suffix == '.pdf':
                return self._extract_pdf(file_path)
            elif suffix in ['.html', '.htm']:
                return self._extract_html(file_path)
            elif suffix in ['.docx', '.doc']:
                return self._extract_docx(file_path)
            elif suffix in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp']:
                return self._extract_image(file_path) if self.ocr_enabled else None
            else:
                # Try as plain text
                return self._extract_txt(file_path)
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return None
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_md(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        return self._extract_txt(file_path)  # Same as txt for now
    
    def _extract_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF with robust pipeline and proper resource cleanup."""
        text_parts = []
        
        # Try PyMuPDF first (fastest)
        doc = None
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            
            for page_num in range(min(len(doc), self.max_pdf_pages)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            if text_parts:
                return '\n\n'.join(text_parts)
                
        except ImportError:
            logger.debug("PyMuPDF not available")
        except Exception as e:
            logger.debug(f"PyMuPDF extraction failed: {e}")
        finally:
            # Always close PDF document, even on error
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass  # Ignore errors during cleanup
        
        # Try pypdfium2
        pdf = None
        try:
            import pypdfium2 as pdfium
            
            pdf = pdfium.PdfDocument(file_path)
            
            for page_num in range(min(len(pdf), self.max_pdf_pages)):
                page = None
                textpage = None
                try:
                    page = pdf[page_num]
                    textpage = page.get_textpage()
                    
                    text = textpage.get_text_bounded()
                    if text.strip():
                        text_parts.append(text)
                finally:
                    # Always close page resources
                    if textpage is not None:
                        try:
                            textpage.close()
                        except Exception:
                            pass
                    if page is not None:
                        try:
                            page.close()
                        except Exception:
                            pass
            
            if text_parts:
                return '\n\n'.join(text_parts)
                
        except ImportError:
            logger.debug("pypdfium2 not available")
        except Exception as e:
            logger.debug(f"pypdfium2 extraction failed: {e}")
        finally:
            # Always close PDF document, even on error
            if pdf is not None:
                try:
                    pdf.close()
                except Exception:
                    pass  # Ignore errors during cleanup
        
        # Try pdfminer as fallback (no resource management needed - handles internally)
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(file_path, maxpages=self.max_pdf_pages)
            if text.strip():
                return text
                
        except ImportError:
            logger.debug("pdfminer not available")
        except Exception as e:
            logger.debug(f"pdfminer extraction failed: {e}")
        
        logger.warning(f"All PDF extraction methods failed for: {file_path}")
        return None
    
    def _extract_html(self, file_path: str) -> Optional[str]:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'lxml')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except ImportError:
            logger.debug("BeautifulSoup not available, using plain text extraction")
            return self._extract_txt(file_path)
        except Exception as e:
            logger.error(f"HTML extraction failed for {file_path}: {e}")
            return None
    
    def _extract_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            return '\n'.join(paragraphs)
            
        except ImportError:
            logger.debug("python-docx not available")
            return None
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return None

    def _extract_image(self, file_path: str) -> Optional[str]:
        """Extract text from image using configured OCR backend (tesseract, paddleocr, easyocr)."""
        extract = get_ocr_extractor(self.ocr_backend)
        text = extract(file_path)
        # Fallback to tesseract if AI backend failed and we're not already using it
        if text is None and self.ocr_backend != "tesseract":
            extract = get_ocr_extractor("tesseract")
            text = extract(file_path)
        return text

    def _chunk_text(self, text: str, file_path, file_id: str) -> List[Chunk]:
        """
        Chunk text into overlapping segments with sentence-aware boundaries.
        
        Uses sentence boundaries (. ! ?) when possible to avoid mid-sentence splits,
        falling back to word-based chunking for text without clear sentence structure.
        """
        max_tokens = self.config["index"]["max_tokens"]
        overlap = self.config["index"]["overlap"]
        use_sentences = self.config["index"].get("sentence_chunking", True)
        
        # Try sentence-aware chunking first (better for prose, docs)
        if use_sentences:
            chunks = self._chunk_by_sentences(text, file_path, file_id, max_tokens, overlap)
            if chunks:
                return chunks
        
        # Fallback: word-based chunking (for code, lists, etc.)
        return self._chunk_by_words(text, file_path, file_id, max_tokens, overlap)
    
    def _chunk_by_sentences(self, text: str, file_path, file_id: str, 
                           max_tokens: int, overlap: int) -> List[Chunk]:
        """Chunk text respecting sentence boundaries."""
        import re
        # Split on sentence boundaries: . ! ? followed by whitespace or end
        sentence_pattern = re.compile(r'(?<=[.!?])\s+|\n\s*\n')
        sentences = [s.strip() for s in sentence_pattern.split(text) if s.strip()]
        
        if not sentences:
            return []
        
        chunks = []
        current_chunk_sentences = []
        current_length = 0
        overlap_sentences = max(1, overlap // 50)  # Overlap ~50 words per sentence
        chunk_idx = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            
            if current_length + sentence_words > max_tokens and current_chunk_sentences:
                # Save current chunk
                chunk_text = ' '.join(current_chunk_sentences)
                chunk = self._create_chunk(chunk_text, file_path, file_id, chunk_idx)
                chunks.append(chunk)
                chunk_idx += 1
                
                # Keep last N sentences for overlap
                overlap_count = min(overlap_sentences, len(current_chunk_sentences))
                current_chunk_sentences = current_chunk_sentences[-overlap_count:]
                current_length = sum(len(s.split()) for s in current_chunk_sentences)
            
            current_chunk_sentences.append(sentence)
            current_length += sentence_words
        
        if current_chunk_sentences:
            chunk_text = ' '.join(current_chunk_sentences)
            chunk = self._create_chunk(chunk_text, file_path, file_id, chunk_idx)
            chunks.append(chunk)
        
        return chunks
    
    def _chunk_by_words(self, text: str, file_path, file_id: str,
                        max_tokens: int, overlap: int) -> List[Chunk]:
        """Fallback: chunk by word count (original behavior)."""
        words = text.split()
        chunks = []
        i = 0
        chunk_idx = 0
        
        while i < len(words):
            chunk_words = words[i:i + max_tokens]
            chunk_text = ' '.join(chunk_words)
            
            if not chunk_text.strip():
                break
            
            chunk = self._create_chunk(chunk_text, file_path, file_id, chunk_idx,
                                       token_start=i, token_end=i + len(chunk_words))
            chunks.append(chunk)
            chunk_idx += 1
            i += max_tokens - overlap
        
        return chunks
    
    def _create_chunk(self, chunk_text: str, file_path, file_id: str, chunk_idx: int,
                      token_start: int = 0, token_end: int = None) -> Chunk:
        """Create a Chunk object with metadata."""
        if token_end is None:
            token_end = token_start + len(chunk_text.split())
        
        return Chunk(
            path=str(file_path),
            file_id=file_id,
            chunk_id=chunk_id(file_id, chunk_idx),
            text=chunk_text,
            token_start=token_start,
            token_end=token_end,
            mtime=int(Path(file_path).stat().st_mtime),
            sha256=hashlib.sha256(chunk_text.encode()).hexdigest(),
            idx=chunk_idx
        )
    
    def _flush_embed_buffer(self):
        """Flush accumulated chunks: embed and upsert to Qdrant."""
        if not self._embed_buffer:
            return
        chunks = self._embed_buffer
        self._embed_buffer = []
        self._embed_and_upsert(chunks)

    def _embed_and_upsert(self, chunks: List[Chunk]):
        """Generate embeddings and upsert to Qdrant (using cached model)."""
        if not chunks:
            return
        try:
            # Use cached model loader (loads once, reuses after)
            model = get_embedding_model()
            if model is None:
                logger.error("Failed to load embedding model")
                return
            
            # Prepare texts
            texts = [chunk.text for chunk in chunks]
            
            # Generate embeddings in large batches (configurable)
            batch_size = self.config["index"].get("embed_batch", 2048)
            embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                batch_embeddings = model.encode(batch_texts, convert_to_tensor=False)
                embeddings.extend(batch_embeddings.tolist())
            
            # Prepare points for Qdrant (payload includes text for dashboard visualization)
            points = []
            for chunk, embedding in zip(chunks, embeddings):
                # Truncate text for payload (Qdrant dashboard display, ~500 chars)
                text_snippet = (chunk.text[:500] + "...") if len(chunk.text) > 500 else chunk.text
                points.append({
                    "id": chunk.chunk_id,
                    "vector": embedding,
                    "payload": {
                        "path": chunk.path,
                        "file_id": chunk.file_id,
                        "chunk_id": chunk.chunk_id,
                        "idx": chunk.idx,
                        "text": text_snippet
                    }
                })
            
            # Upsert to Qdrant in large batches
            self.qdrant.upsert_vectors(points)
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
    
    def _should_exclude(self, path: str) -> bool:
        """Check if path should be excluded."""
        from fnmatch import fnmatch
        
        for pattern in self.exclude_patterns:
            if fnmatch(path, pattern):
                return True
        return False
    
    def _get_existing_sha256(self, file_id: str) -> Optional[str]:
        """Get existing SHA256 for file if it exists."""
        try:
            cursor = self.catalog.conn.execute(
                "SELECT sha256 FROM files WHERE file_id = ?", (file_id,)
            )
            row = cursor.fetchone()
            return row["sha256"] if row else None
        except Exception:
            return None
    
    def _load_frontier(self) -> FrontierState:
        """Load frontier state from disk."""
        try:
            if self.frontier_path.exists():
                with open(self.frontier_path, 'r') as f:
                    data = json.load(f)
                
                # Convert seen dict keys back to strings (JSON keys are strings)
                seen = {k: tuple(v) for k, v in data.get("seen", {}).items()}
                
                return FrontierState(
                    queue=data.get("queue", []),
                    seen=seen,
                    processed_files=data.get("processed_files", 0),
                    processed_dirs=data.get("processed_dirs", 0),
                    errors=data.get("errors", [])
                )
        except Exception as e:
            logger.warning(f"Failed to load frontier: {e}")
        
        return FrontierState(queue=[], seen={})
    
    def _save_frontier(self, frontier: FrontierState):
        """Save frontier state to disk."""
        try:
            self.frontier_path.parent.mkdir(parents=True, exist_ok=True)
            
            data = {
                "queue": frontier.queue,
                "seen": {str(k): list(v) for k, v in frontier.seen.items()},
                "processed_files": frontier.processed_files,
                "processed_dirs": frontier.processed_dirs,
                "errors": frontier.errors
            }
            
            with open(self.frontier_path, 'w') as f:
                json.dump(data, f, indent=2)
                
        except Exception as e:
            logger.error(f"Failed to save frontier: {e}")


def run_bfs_slice(roots: List[str], **kwargs) -> IndexStats:
    """Run one BFS slice with given parameters."""
    from .validation import sanitize_file_path, validate_chunk_params

    # Validate roots
    valid_roots = []
    for r in roots:
        resolved, err = sanitize_file_path(r, must_exist=True)
        if err:
            logger.warning(f"Skipping invalid root {r}: {err}")
            continue
        valid_roots.append(str(resolved))
    if not valid_roots:
        raise ValueError("No valid index roots provided")

    config = get_config()

    # Validate chunk params from config
    max_tokens = config["index"].get("max_tokens", 1200)
    overlap = config["index"].get("overlap", 80)
    valid, err = validate_chunk_params(max_tokens, overlap)
    if not valid:
        logger.warning(f"Chunk params: {err}, using defaults")
        config["index"]["max_tokens"] = min(max(1200, max_tokens), 10000)
        config["index"]["overlap"] = min(max(0, overlap), max_tokens - 1)

    # Override config with kwargs
    for key, value in kwargs.items():
        if key in config:
            config[key] = value
        elif key in config.get("index", {}):
            config["index"][key] = value
    
    indexer = BFSIndexer(config)

    start_time = time.time()
    stats = indexer.run_bfs_slice(valid_roots, kwargs.get("max_items", 1000))
    stats.duration_seconds = time.time() - start_time
    
    logger.info(f"BFS slice completed: {stats.files_processed} files, {stats.chunks_created} chunks, {stats.duration_seconds:.2f}s")
    
    return stats


def run_complete_index(roots: List[str], **kwargs) -> IndexStats:
    """Run complete indexing of all files in roots."""
    from .validation import sanitize_file_path

    # Validate roots
    valid_roots = []
    for r in roots:
        resolved, err = sanitize_file_path(r, must_exist=True)
        if err:
            logger.warning(f"Skipping invalid root {r}: {err}")
            continue
        valid_roots.append(str(resolved))
    if not valid_roots:
        raise ValueError("No valid index roots provided")

    config = get_config()

    # Override config with kwargs
    for key, value in kwargs.items():
        if key in config:
            config[key] = value
        elif key in config.get("index", {}):
            config["index"][key] = value

    indexer = BFSIndexer(config)

    start_time = time.time()

    # Run BFS slices until all files are processed
    total_stats = IndexStats()
    max_items_per_slice = kwargs.get("max_items_per_slice", 1000)

    # Clear frontier to start fresh
    frontier_path = Path(config["paths"]["frontier"])
    if frontier_path.exists():
        frontier_path.unlink()
        logger.info("Cleared existing frontier for fresh start")

    while True:
        # Run one slice
        slice_stats = indexer.run_bfs_slice(valid_roots, max_items_per_slice)
        
        # Accumulate stats
        total_stats.files_processed += slice_stats.files_processed
        total_stats.chunks_created += slice_stats.chunks_created
        total_stats.files_skipped += slice_stats.files_skipped
        total_stats.errors += slice_stats.errors
        
        # Load current frontier to check if there are more items to process
        frontier = indexer._load_frontier()
        
        if not frontier.queue:
            logger.info("No more items to process. Indexing complete.")
            break
        
        # If no files were processed AND no items were added to frontier, we're done
        if slice_stats.files_processed == 0 and len(frontier.queue) == 0:
            logger.info("No files processed and no items in queue. Indexing complete.")
            break
        
        logger.info(f"Processed {slice_stats.files_processed} files, {len(frontier.queue)} items remaining in queue")
    
    total_stats.duration_seconds = time.time() - start_time
    
    logger.info(f"Complete indexing finished: {total_stats.files_processed} files, {total_stats.chunks_created} chunks, {total_stats.duration_seconds:.2f}s")
    
    return total_stats


if __name__ == "__main__":
    # Test the indexer
    import sys

    default_root = str(Path.home() / "Documents")
    roots = [sys.argv[1]] if len(sys.argv) > 1 else [default_root]
    
    print(f"Testing BFS indexer with roots: {roots}")
    
    stats = run_bfs_slice(roots, max_items=10)
    
    print(f"Results:")
    print(f"  Files processed: {stats.files_processed}")
    print(f"  Chunks created: {stats.chunks_created}")
    print(f"  Files skipped: {stats.files_skipped}")
    print(f"  Errors: {stats.errors}")
    print(f"  Duration: {stats.duration_seconds:.2f}s")
